# Legal Template Filler – Streamlit App
# -------------------------------------------------------------
# This single-file Streamlit app lets you:
# 1) Upload a .docx legal template
# 2) Automatically detect placeholders ({{like_this}}, [[LIKE_THIS]], or [ALL_CAPS_TOKEN])
# 3) Chat to fill each placeholder conversationally
# 4) Preview the filled document (HTML) and download the final .docx
#
# Quick Start (Local):
#   1) Save this file as app.py
#   2) Create a virtual environment (optional):
#        python -m venv .venv && source .venv/bin/activate  # (Windows: .venv\\Scripts\\activate)
#   3) Install deps:
#        pip install streamlit python-docx mammoth regex
#   4) Run:
#        streamlit run app.py
#   5) Open the local URL shown in your terminal.
#
# Free Hosting (Public URL):
#   * Streamlit Community Cloud (https://streamlit.io/cloud)
#     - Push this file to a public GitHub repo
#     - Create a new Streamlit app from that repo (main file: app.py)
#     - The app gets a public URL automatically
#
# Notes:
# - We use three placeholder styles by default:
#     {{client_name}}  |  [[EFFECTIVE_DATE]]  |  [GOVERNING_LAW]
# - You can change or extend patterns in PLACEHOLDER_PATTERNS below.
# - We show a highlighted preview using Mammoth (docx->HTML) + simple regex markup.
# - Replacement traverses paragraphs, tables, headers/footers for robust coverage.
#
# -------------------------------------------------------------

import io
import re
import regex as rxx
from dataclasses import dataclass
from typing import Dict, List, Set, Tuple

import streamlit as st
from docx import Document
import mammoth

# -----------------------------
# Configuration
# -----------------------------

# Regex patterns to detect placeholders in text.
# Order matters (more specific first). Each pattern MUST capture the placeholder name in group 1.
PLACEHOLDER_PATTERNS: List[Tuple[str, str]] = [
    (r"\{\{\s*([A-Za-z0-9_\- ]+)\s*\}\}", "double_curly"),        # {{placeholder}}
    (r"\[\[\s*([A-Za-z0-9_\- ]+)\s*\]\]", "double_square"),        # [[PLACEHOLDER]] / [[Client Name]]
    (r"\[(_{3,})\]", "underscored_blank"),                               # [__________] 3+ underscores
    (r"\[([A-Za-z][A-Za-z0-9 _\-/&\.,]{1,})\]", "human_readable"),     # [Company Name], [Governing law]
    (r"\[([A-Z][A-Z0-9_\-]{2,})\]", "caps_brackets"),                  # [ALL_CAPS_TOKEN]
]

HIGHLIGHT_COLOR = "#fff7c2"  # pale yellow


# -----------------------------
# Utilities
# -----------------------------

def is_underscore_key(key: str) -> bool:
    return bool(re.fullmatch(r"_+", key or ""))


def _read_docx_bytes(file_bytes: bytes) -> Document:
    bio = io.BytesIO(file_bytes)
    return Document(bio)


def _extract_text_runs(doc: Document):
    """(Deprecated) Kept for reference; run-level ops miss split placeholders."""
    for para in doc.paragraphs:
        for run in para.runs:
            yield run, "paragraph"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        yield run, "table"

def _extract_headers_footers(doc: Document):
    parts = []
    for section in doc.sections:
        if section.header:
            parts.append(section.header)
        if section.footer:
            parts.append(section.footer)
    return parts


def _iter_runs_in_header_footer(hf) -> List:
    for para in hf.paragraphs:
        for run in para.runs:
            yield run

# NEW: paragraph helpers that work across run boundaries

def _iter_all_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for hf in _extract_headers_footers(doc):
        for para in hf.paragraphs:
            yield para


def _iter_runs_in_header_footer(hf) -> List:
    for para in hf.paragraphs:
        for run in para.runs:
            yield run


def find_placeholders_in_text(text: str) -> List[str]:
    found = []
    for pattern, _tag in PLACEHOLDER_PATTERNS:
        for m in re.finditer(pattern, text):
            name = m.group(1)
            if name:
                found.append(name)
    return found


def extract_placeholders(doc: Document) -> Set[str]:
    """Find placeholders by scanning full paragraph text (handles split runs)."""
    placeholders: Set[str] = set()
    for para in _iter_all_paragraphs(doc):
        text = "".join(run.text for run in para.runs)
        placeholders.update(find_placeholders_in_text(text))
    return placeholders


def replace_placeholders_in_text(text: str, mapping: Dict[str, str]) -> str:
    # Build a single pass replacer for all patterns.
    def repl_factory(pattern):
        regex = re.compile(pattern)
        def _repl(m):
            key = m.group(1)
            return mapping.get(key, m.group(0))  # leave as is if missing
        return regex.sub, _repl

    for pattern, _tag in PLACEHOLDER_PATTERNS:
        sub_fn, repl = repl_factory(pattern)
        text = sub_fn(repl, text)
    return text


def fill_document(doc: Document, mapping: Dict[str, str]) -> Document:
    """Replace placeholders at the paragraph level to survive DOCX run splits.
    We collapse runs per paragraph by setting the first run's text to the
    replaced string and blanking the remaining runs (formatting minimalism).
    """
    for para in _iter_all_paragraphs(doc):
        # Aggregate text across runs
        original = "".join(run.text for run in para.runs)
        if any(ch in original for ch in ['{', '}', '[']):
            replaced = replace_placeholders_in_text(original, mapping)
            # Write back
            if para.runs:
                para.runs[0].text = replaced
                for r in para.runs[1:]:
                    r.text = ""
    return doc


def docx_to_html(doc_bytes: bytes, highlight_mapping: Dict[str, str] = None) -> str:
    """Convert docx -> HTML and optionally highlight unresolved placeholders.
    For highlighting, we simply wrap literal token forms like [Key], [[Key]], {{Key}}.
    This works even when keys contain spaces/underscores.
    """
    with io.BytesIO(doc_bytes) as f:
        result = mammoth.convert_to_html(f)
        html = result.value

    if highlight_mapping:  # highlight unresolved
        for key in highlight_mapping.keys():
            # Prepare all literal token renderings we support
            forms = [f"[{key}]", f"[[{key}]]", f"{{{{{key}}}}}"]
            for token in forms:
                pattern = re.escape(token)
                html = re.sub(pattern, lambda m: f'<mark style="background:{HIGHLIGHT_COLOR}">{m.group(0)}</mark>', html)
    return html


def guess_question_from_key(key: str) -> str:
    # Friendly prompt from a placeholder key (supports underscores-only tokens)
    if is_underscore_key(key):
        # Use a stable index label from session state
        idx = None
        if "underscore_indices" in st.session_state:
            idx = st.session_state.underscore_indices.get(key)
        label = f"Blank {idx}" if idx is not None else "Blank"
        return f"Please provide value for {label}:"
    pretty = key.replace("_", " ").replace("-", " ")
    pretty = pretty.strip()
    if pretty:
        pretty = pretty[0].upper() + pretty[1:]
    else:
        pretty = "value"
    return f"Please provide {pretty}:"


# -----------------------------
# Streamlit UI
# -----------------------------

st.set_page_config(page_title="Legal Template Filler", page_icon="📄", layout="wide")

# App Title
st.title("📄 Legal Template Filler")
st.caption("Upload a .docx, chat to fill placeholders, preview, and download.")

# Session state
if "placeholders" not in st.session_state:
    st.session_state.placeholders = set()
if "values_map" not in st.session_state:
    st.session_state.values_map = {}
    st.session_state.values_map = {}
if "pending_keys" not in st.session_state:
    st.session_state.pending_keys = []
if "raw_docx" not in st.session_state:
    st.session_state.raw_docx = None
if "messages" not in st.session_state:
    st.session_state.messages = []

# Sidebar: instructions
with st.sidebar:
    st.header("How it works")
    st.markdown(
        """
        1. **Upload** your .docx legal template
        2. We **detect placeholders** (e.g., `{{client_name}}`, `[[EFFECTIVE_DATE]]`, `[GOVERNING_LAW]`, `[Company Name]`, `[__________]`, `[ALL_CAPS_TOKEN]`)
        3. Use the **chat** to fill them in
        4. **Preview** & **Download** the completed document
        """
    )

# File uploader
uploaded = st.file_uploader("Upload a .docx template", type=["docx"], help="The app scans for placeholders and guides you to fill them.")

# Optional sample
with st.expander("Or load a sample NDA template"):
    SAMPLE = st.checkbox("Load sample", value=False, help="Use a tiny built-in NDA-style template for demo.")

if SAMPLE:
    from docx import Document as Doc
    buf = io.BytesIO()
    d = Doc()
    d.add_heading("Mutual Non-Disclosure Agreement", 0)
    p = d.add_paragraph(
        "This Mutual Non-Disclosure Agreement (the \"Agreement\") is made as of [[EFFECTIVE_DATE]] between {{company_name}}, a [STATE_OF_INCORPORATION] corporation, and {{counterparty_name}}. "
    )
    d.add_paragraph(
        "The parties agree that confidential information disclosed under this Agreement shall be used solely for the purpose of [[PURPOSE]] and governed by the laws of [GOVERNING_LAW]."
    )
    d.add_paragraph("Authorized signatory for {{company_name}}: [[SIGNATORY_NAME]].")
    d.save(buf)
    st.session_state.raw_docx = buf.getvalue()
elif uploaded is not None:
    st.session_state.raw_docx = uploaded.read()

# If we have a document, process it
if st.session_state.raw_docx:
    doc = _read_docx_bytes(st.session_state.raw_docx)
    placeholders = extract_placeholders(doc)
    # Build stable labels for underscore placeholders
    if "underscore_indices" not in st.session_state:
        st.session_state.underscore_indices = {}
    next_idx = 1 + len(st.session_state.underscore_indices)
    for k in sorted(placeholders):
        if is_underscore_key(k) and k not in st.session_state.underscore_indices:
            st.session_state.underscore_indices[k] = next_idx
            next_idx += 1
    st.session_state.placeholders = placeholders
    # Initialize pending keys
    st.session_state.pending_keys = [k for k in sorted(placeholders) if k not in st.session_state.values_map]

    # Layout: left chat, right preview/panels
    left, right = st.columns([0.6, 0.4])

    # ---------------- Chat Pane -----------------
    with left:
        st.subheader("Chat to fill placeholders")

        # Initial system message
        if not st.session_state.messages:
            if placeholders:
                first_key = st.session_state.pending_keys[0] if st.session_state.pending_keys else None
                prompt = guess_question_from_key(first_key) if first_key else "No placeholders found."
                st.session_state.messages.append({"role": "assistant", "content": f"I found {len(placeholders)} placeholders. {prompt}"})
            else:
                st.session_state.messages.append({"role": "assistant", "content": "I didn't find any placeholders. You can still preview or upload another template."})

        # Render past messages
        for m in st.session_state.messages:
            with st.chat_message(m["role"]):
                st.markdown(m["content"])

        # Input box
        user_input = st.chat_input("Type your answer or ask a question…")

        def advance_to_next_key():
            # Drop keys already filled
            st.session_state.pending_keys = [k for k in st.session_state.pending_keys if k not in st.session_state.values_map]
            if st.session_state.pending_keys:
                nxt = st.session_state.pending_keys[0]
                q = guess_question_from_key(nxt)
                st.session_state.messages.append({"role": "assistant", "content": q})
            else:
                st.session_state.messages.append({"role": "assistant", "content": "All placeholders are filled ✅. You can preview and download now."})

        if user_input is not None:
            # Decide whether user is answering current placeholder or chatting
            st.session_state.messages.append({"role": "user", "content": user_input})

            if st.session_state.pending_keys:
                current_key = st.session_state.pending_keys[0]
                # Very simple heuristic: treat any user input as the value for the current placeholder
                st.session_state.values_map[current_key] = user_input.strip()
                st.session_state.messages.append({"role": "assistant", "content": f"Got it. **{current_key}** set to: `{user_input.strip()}`"})
                advance_to_next_key()
            else:
                # No placeholders pending – answer simple help intents
                if user_input.strip().lower() in {"reset", "restart", ":reset"}:
                    st.session_state.values_map = {}
                    st.session_state.pending_keys = sorted(st.session_state.placeholders)
                    st.session_state.messages.append({"role": "assistant", "content": "Reset complete. Let's start again."})
                    advance_to_next_key()
                else:
                    st.session_state.messages.append({"role": "assistant", "content": "Everything looks complete. You can still change values from the sidebar."})

        # Manual edit form
        with st.expander("Edit values manually"):
            if placeholders:
                for k in sorted(placeholders):
                    st.session_state.values_map[k] = st.text_input(k, value=st.session_state.values_map.get(k, ""), key=f"manual_{k}")
                if st.button("Apply changes", key="apply_changes"):
                    st.session_state.messages.append({"role": "assistant", "content": "Values updated."})

    # ---------------- Right Pane -----------------
    with right:
        st.subheader("Preview & Download")

        # Show detected placeholders
        if placeholders:
            st.markdown("**Detected placeholders:**")
            st.write(sorted(placeholders))
        else:
            st.info("No placeholders detected.")

        # Compose unresolved keys for highlight
        unresolved = {k: "" for k in placeholders if not st.session_state.values_map.get(k)}

        # Produce a filled .docx in-memory
        working = _read_docx_bytes(st.session_state.raw_docx)
        filled = fill_document(working, st.session_state.values_map)
        out_buf = io.BytesIO()
        filled.save(out_buf)
        out_bytes = out_buf.getvalue()

        # HTML preview with unresolved highlights
        html = docx_to_html(out_bytes, highlight_mapping=unresolved)
        st.markdown("**Document preview:**")
        st.components.v1.html(
            f"""
            <div style='border:1px solid #ddd;border-radius:8px;padding:16px;max-height:60vh;overflow:auto;background:#fff'>
                {html}
            </div>
            """,
            height=500,
            scrolling=True,
        )

        # Download button
        st.download_button(
            label="⬇️ Download .docx",
            data=out_bytes,
            file_name="completed_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.caption("Tip: Any placeholders still visible in the preview will be highlighted.")

else:
    st.info("Upload a .docx or load the sample from the expander to begin.")
